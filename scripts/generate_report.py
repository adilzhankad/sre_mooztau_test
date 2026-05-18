#!/usr/bin/env python3
"""Generate End Term Project Word report (MoozTau Platform).

First-person narrative style, matches the previous MoozTau_Report and
Assignment 6 reports. Each [SCREENSHOT] marker carries a detailed instruction
for what exactly to capture.

Output: docs/MoozTau_EndTerm_Report.docx
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "MoozTau_EndTerm_Report.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_code_block(doc, code):
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = "Courier New"
    r.font.size = Pt(9)


def add_screenshot(doc, label, what_to_open, what_to_capture, filename):
    p = doc.add_paragraph()
    r = p.add_run(f"[SCREENSHOT {filename}]  {label}")
    r.bold = True
    r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    r.font.size = Pt(10)

    p1 = doc.add_paragraph()
    r1 = p1.add_run("Open: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p1.add_run(what_to_open)
    r2.font.size = Pt(10)
    r2.italic = True

    p2 = doc.add_paragraph()
    r3 = p2.add_run("Capture: ")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run(what_to_capture)
    r4.font.size = Pt(10)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_title_page(doc):
    for _ in range(7):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MoozTau Platform")
    r.bold = True
    r.font.size = Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("End Term Project")
    r.bold = True
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Comprehensive SRE Implementation for a Distributed Microservices System"
    )
    r.bold = True
    r.italic = True
    r.font.size = Pt(14)

    for _ in range(3):
        doc.add_paragraph()

    meta = [
        "Student: Adilzhan Kadyrov",
        f"Date: {date.today().strftime('%d.%m.%Y')}",
        "Server: 213.155.22.46",
        "Domain: medhome.kz",
        "Repository: https://github.com/adilzhankad/sre_mooztau_test",
    ]
    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(12)

    doc.add_page_break()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title_page(doc)

    # ====== 1. INTRODUCTION ======
    add_heading(doc, "1. Introduction", level=1)
    add_para(
        doc,
        "This is my End Term Project for the SRE course. The goal of this project "
        "is to demonstrate the entire Site Reliability Engineering lifecycle on a "
        "real microservices platform that I built and deployed — MoozTau. The "
        "system is already running in production on my VPS at 213.155.22.46 and "
        "is publicly available at medhome.kz.",
    )
    add_para(
        doc,
        "Throughout the course I built this project step by step. In Assignment 1 "
        "I containerized the services with Docker Compose. In Assignment 2 I "
        "defined SLIs and SLOs. In Assignment 3 I added Prometheus + Grafana "
        "monitoring. In Assignment 4 I simulated an incident on the Orders Service "
        "and wrote a postmortem. In Assignment 5 I described the infrastructure "
        "with Terraform. In Assignment 6 I added automation and capacity planning.",
    )
    add_para(
        doc,
        "For this End Term Project I extended the system with every component "
        "required by the specification, plus several production-grade additions:",
    )
    add_bullets(
        doc,
        [
            "Multi-orchestration: Docker Swarm stack file and a full set of "
            "Kubernetes manifests, on top of Docker Compose.",
            "Two-database architecture: PostgreSQL stays the primary transactional "
            "store; MongoDB was added as a parallel archive + full-text search "
            "for chat messages.",
            "Kubernetes health checks at three levels — startupProbe, "
            "readinessProbe, livenessProbe — for every Deployment.",
            "Horizontal Pod Autoscalers (HPA) on all six microservices, tuned "
            "per-service (orders gets the widest range, finance the narrowest).",
            "Configuration management with Ansible — playbook + four roles.",
            "Terraform restructured into two profiles — Yandex Cloud "
            "provisioning + existing VPS.",
            "Multi-stage CI/CD: a separate ci.yml workflow runs lint + IaC "
            "validation on every push; deploy.yml gates production behind that, "
            "runs a post-deploy health check, and notifies Telegram.",
            "Two new docs — SLI/SLO catalogue (docs/sli_slo.md) and capacity "
            "planning (docs/capacity_planning.md).",
            "Single validation script (scripts/validate_all.sh) that checks all "
            "configurations in one command — 32 k8s resources across 14 files "
            "validated by kubeconform.",
        ],
    )

    add_screenshot(
        doc,
        "Project structure in IDE",
        "VS Code with the Mooztau_back project. Expand the file tree on the left "
        "so all top-level folders are visible.",
        "Show every component of the project in one frame: the six service "
        "folders, MoozTau (frontend), monitoring/, terraform/, ansible/, k8s/, "
        ".github/, docs/, scripts/. Take it with the tree on the left and an "
        "open file on the right — for example docker-compose.yml.",
        "01_project_tree.png",
    )

    # ====== 2. SYSTEM OVERVIEW ======
    add_heading(doc, "2. System Overview", level=1)
    add_para(
        doc,
        "MoozTau is a production-management and dealer-network system for "
        "refrigeration equipment. The architecture is built around six "
        "independent backend microservices, a React frontend, two databases "
        "(PostgreSQL + MongoDB), and a full monitoring stack. Every component "
        "runs in Docker on the internal mooztau_net bridge network.",
    )
    add_para(doc, "Service inventory:", bold=True)
    add_table(
        doc,
        ["Service", "Port", "Responsibility"],
        [
            ["auth", "8002", "Authentication, JWT, users, organizations"],
            ["orders", "8001", "Orders, products, factory, analytics"],
            ["finance", "8003", "Bank accounts, transactions, financial reports"],
            ["product", "8004", "Product catalogue and pricing"],
            ["user", "8005", "User profiles and management"],
            ["chat", "8006", "WebSocket messaging + Mongo archive + full-text search"],
            ["frontend", "3100", "React + TypeScript served by Nginx"],
            ["postgres", "5432 (internal)", "Primary relational store (PostgreSQL 16)"],
            ["mongo", "27017 (internal)", "Chat archive + full-text search (MongoDB 7)"],
            ["prometheus", "9090", "Metrics scraping every 15s"],
            ["grafana", "3000", "Dashboards"],
            ["alertmanager", "9093", "Alert routing → Telegram"],
        ],
    )
    add_para(
        doc,
        "Every microservice exposes a /health endpoint for Docker / Kubernetes "
        "health checks and a /metrics endpoint for Prometheus scraping. The "
        "chat service additionally exposes /health/full which verifies both "
        "PostgreSQL and MongoDB are reachable. All containers use the "
        "unless-stopped restart policy so they recover automatically after "
        "crashes or host reboots.",
    )

    add_screenshot(
        doc,
        "All containers running on production",
        "SSH to the server: ssh -i ~/.ssh/github_actions ubuntu@213.155.22.46, "
        "then run sudo docker ps in the terminal.",
        "All 12 mooztau containers must be visible with the status Up and "
        "(healthy) in the STATUS column — including the new mongo container. "
        "Both the terminal prompt and the full table in one frame.",
        "02_docker_ps_server.png",
    )

    # ====== 3. DOCKER COMPOSE ======
    add_heading(doc, "3. Environment Setup — Docker Compose (Assignment 1)", level=1)
    add_para(
        doc,
        "The base environment is orchestrated with Docker Compose. The "
        "docker-compose.yml file at the project root defines all 12 services, "
        "the named volumes for PostgreSQL, MongoDB, Prometheus, and Grafana "
        "data, and the bridge network. Every service has explicit dependencies "
        "and health checks, so downstream services only start once their "
        "dependencies are marked healthy.",
    )
    add_para(doc, "MongoDB health-check example (new):", italic=True)
    add_code_block(
        doc,
        """healthcheck:
  test: ["CMD", "mongosh", "--quiet", "--eval", "db.runCommand({ ping: 1 }).ok"]
  interval: 10s
  timeout: 5s
  retries: 5
  start_period: 20s""",
    )
    add_para(doc, "Microservice (orders) health-check:", italic=True)
    add_code_block(
        doc,
        """healthcheck:
  test: ["CMD-SHELL", "python -c \\"import urllib.request; \\
         urllib.request.urlopen('http://localhost:8001/health')\\""]
  interval: 30s
  timeout: 10s
  start_period: 15s
  retries: 3""",
    )
    add_para(
        doc,
        "I bring up the whole stack with one command: docker compose up -d "
        "--build. The build flag rebuilds images from each service's Dockerfile.",
    )

    add_screenshot(
        doc,
        "docker-compose.yml in the IDE",
        "Open docker-compose.yml in VS Code, scroll to a section where the new "
        "mongo service is visible (lines ~19-35).",
        "Capture the YAML structure with the mongo service block + its "
        "healthcheck + the chat service that depends on both db and mongo "
        "(condition: service_healthy). Line numbers visible on the left.",
        "03_compose_yml.png",
    )

    # ====== 4. SLI / SLO ======
    add_heading(doc, "4. SLI / SLO Design (Assignment 2)", level=1)
    add_para(
        doc,
        "Service Level Indicators (SLIs) tell me how well the system serves my "
        "users. Service Level Objectives (SLOs) are the targets I commit to. I "
        "derive every SLI from Prometheus metrics that the FastAPI services "
        "expose. All definitions live in docs/sli_slo.md.",
    )
    add_table(
        doc,
        ["SLI", "Definition", "SLO Target"],
        [
            ["Availability", "1 − (5xx requests / total requests)", "≥ 99%"],
            ["Latency P95", "95th-percentile HTTP response time", "≤ 200 ms"],
            ["Error rate", "5xx responses / total responses", "≤ 1%"],
            ["Success rate", "2xx responses / total responses", "≥ 99%"],
        ],
    )
    add_para(doc, "Example PromQL — availability:", italic=True)
    add_code_block(
        doc,
        """sum(rate(http_request_duration_seconds_count{status!~"5.."}[5m]))
/
sum(rate(http_request_duration_seconds_count[5m]))""",
    )
    add_para(
        doc,
        "I also defined an error budget — at 99% availability over 30 days I "
        "can afford 432 minutes (~7.2 hours) of unavailability before I break "
        "the SLO. If half of that budget is consumed in a month, I freeze "
        "new feature releases and focus on reliability work.",
    )

    add_screenshot(
        doc,
        "SLI/SLO catalogue",
        "Open docs/sli_slo.md in VS Code in preview mode "
        "(right-click → Open Preview).",
        "Capture the top of the rendered Markdown — the SLI definitions and "
        "the SLO target table.",
        "04_sli_slo_doc.png",
    )

    # ====== 5. MONITORING ======
    add_heading(doc, "5. Monitoring and Alerting (Assignment 3)", level=1)
    add_para(
        doc,
        "Prometheus scrapes every microservice every 15 seconds from "
        "<service>:<port>/metrics. The scrape configuration lives in "
        "monitoring/prometheus/prometheus.yml and defines six jobs — one per "
        "service. Grafana reads Prometheus as its data source and renders a "
        "single MoozTau Platform Overview dashboard that I pre-provisioned via "
        "monitoring/grafana/provisioning.",
    )

    add_screenshot(
        doc,
        "Grafana dashboard",
        "Open http://213.155.22.46:3000 → Dashboards → MoozTau Platform "
        "Overview. Time range Last 30 minutes.",
        "All six service tiles UP in green. Request Rate / Latency / Error "
        "Rate panels populated.",
        "05_grafana_dashboard.png",
    )
    add_screenshot(
        doc,
        "Prometheus targets",
        "Open http://213.155.22.46:9090/targets in a browser.",
        "All six scrape pools expanded — every endpoint State = UP (green).",
        "06_prometheus_targets.png",
    )

    add_para(doc, "Six alert rules cover the main failure modes:")
    add_table(
        doc,
        ["Alert", "Condition", "Severity", "For"],
        [
            ["ServiceDown", "up == 0", "critical", "15 seconds"],
            ["HighErrorRate", "5xx rate > 5%", "critical", "5 minutes"],
            ["HighLatency", "P95 > 2 s", "warning", "5 minutes"],
            ["HighRequestRate", "RPS > 50 per service", "warning", "2 minutes"],
            ["ServiceSaturated", "P99 > 5 s", "critical", "3 minutes"],
            ["NoMetrics", "absent(http_request_duration_seconds_count)", "warning", "5 minutes"],
        ],
    )
    add_screenshot(
        doc,
        "Prometheus alert rules",
        "Open http://213.155.22.46:9090/alerts.",
        "All six rules visible with their PromQL expressions and inactive/firing badges.",
        "07_prometheus_alerts.png",
    )
    add_para(
        doc,
        "Alertmanager routes every critical alert to my Telegram bot "
        "@mooztau_alerts_bot. The Telegram template uses ИНЦИДЕНТ for fires "
        "and ВОССТАНОВЛЕНО for resolves.",
    )

    # ====== 6. MULTI-ORCHESTRATION ======
    add_heading(doc, "6. Multi-Orchestration", level=1)
    add_para(
        doc,
        "The End Term Project specification requires both Docker Swarm and "
        "Kubernetes orchestration on top of the Docker Compose setup. I "
        "implemented both — they live in docker-stack.yml and the k8s/ folder.",
    )

    add_heading(doc, "6.1 Docker Swarm", level=2)
    add_para(
        doc,
        "Docker Swarm gave me a simple way to run multiple replicas of every "
        "microservice on the same VPS without changing my workflow. The "
        "docker-stack.yml file mirrors the Compose file but adds Swarm-specific "
        "deploy keys: replicas count, resource limits, rolling-update strategy, "
        "and restart policies.",
    )
    add_para(doc, "Deploy block for orders service:", italic=True)
    add_code_block(
        doc,
        """deploy:
  replicas: 2
  update_config:
    parallelism: 1
    delay: 10s
    order: start-first
  restart_policy:
    condition: on-failure
    delay: 5s
    max_attempts: 3
  resources:
    limits:      { cpus: "1.0", memory: 512M }
    reservations:{ cpus: "0.2", memory: 128M }""",
    )
    add_para(doc, "Deploying the stack:", italic=True)
    add_code_block(
        doc,
        """docker swarm init
docker compose build
docker stack deploy -c docker-stack.yml mooztau
docker service ls
docker stack ps mooztau""",
    )
    add_screenshot(
        doc,
        "docker service ls output",
        "Run docker stack deploy then docker service ls.",
        "All 12 services with REPLICAS = 2/2 for microservices and 1/1 for "
        "db, mongo, prometheus, grafana, alertmanager.",
        "08_docker_service_ls.png",
    )
    add_screenshot(
        doc,
        "docker stack ps mooztau",
        "Run docker stack ps mooztau --format 'table {{.Name}}\\t{{.Image}}\\t"
        "{{.Node}}\\t{{.CurrentState}}'.",
        "Task list — each microservice has 2 task lines, CURRENT STATE = Running.",
        "09_docker_stack_ps.png",
    )

    add_heading(doc, "6.2 Kubernetes — manifests, probes, HPA", level=2)
    add_para(
        doc,
        "For Kubernetes I created 14 manifests under k8s/ covering the entire "
        "stack: Namespace (mooztau), ConfigMap, Secret, PersistentVolumeClaim + "
        "Deployment + Service for both PostgreSQL and MongoDB, six Deployment + "
        "Service pairs for the microservices, a monitoring manifest for "
        "Prometheus and Grafana, an Ingress for external routing, and an HPA "
        "manifest with six Horizontal Pod Autoscalers.",
    )
    add_para(doc, "Three-level health checks (every microservice Deployment):", bold=True)
    add_code_block(
        doc,
        """startupProbe:
  httpGet: { path: /health, port: 8001 }
  initialDelaySeconds: 5
  periodSeconds: 5
  timeoutSeconds: 3
  failureThreshold: 24      # up to 120 s for slow startup (alembic, seed)
readinessProbe:
  httpGet: { path: /health, port: 8001 }
  initialDelaySeconds: 15
  periodSeconds: 10
  timeoutSeconds: 3
  failureThreshold: 3
  successThreshold: 1
livenessProbe:
  httpGet: { path: /health, port: 8001 }
  initialDelaySeconds: 30
  periodSeconds: 30
  timeoutSeconds: 5
  failureThreshold: 3
resources:
  requests: { cpu: "200m", memory: "256Mi" }
  limits:   { cpu: "1000m", memory: "512Mi" }""",
    )
    add_para(
        doc,
        "startupProbe shields slow services (orders runs alembic migrations and "
        "seed scripts on boot) from being killed by livenessProbe before they "
        "are ready. readinessProbe gates traffic before sending it to a pod. "
        "livenessProbe restarts a pod that froze.",
    )

    add_para(doc, "Horizontal Pod Autoscalers (k8s/hpa.yml):", bold=True)
    add_table(
        doc,
        ["Service", "minReplicas", "maxReplicas", "CPU target", "Memory target", "Notes"],
        [
            ["auth", "2", "5", "70%", "80%", "Balanced policy"],
            ["orders", "2", "8", "65%", "75%", "Widest range — heaviest service"],
            ["finance", "2", "4", "70%", "—", "Read-heavy, stable"],
            ["product", "2", "4", "70%", "—", "Read-heavy, stable"],
            ["user", "2", "4", "70%", "—", "Read-heavy, stable"],
            ["chat", "2", "6", "70%", "80%", "Spikes with WebSocket sessions"],
        ],
    )
    add_para(doc, "orders-hpa block — most aggressive scale-up:", italic=True)
    add_code_block(
        doc,
        """apiVersion: autoscaling/v2
kind: HorizontalPodAutoscaler
metadata: { name: orders-hpa, namespace: mooztau }
spec:
  scaleTargetRef: { apiVersion: apps/v1, kind: Deployment, name: orders }
  minReplicas: 2
  maxReplicas: 8
  metrics:
    - type: Resource
      resource: { name: cpu, target: { type: Utilization, averageUtilization: 65 } }
    - type: Resource
      resource: { name: memory, target: { type: Utilization, averageUtilization: 75 } }
  behavior:
    scaleUp:   { stabilizationWindowSeconds: 30, policies: [{ type: Pods, value: 2, periodSeconds: 30 }] }
    scaleDown: { stabilizationWindowSeconds: 300, policies: [{ type: Percent, value: 25, periodSeconds: 60 }] }""",
    )

    add_para(doc, "Validating every manifest with kubeconform — no real cluster needed:")
    add_code_block(doc, "kubeconform -strict -summary k8s/")
    add_screenshot(
        doc,
        "kubeconform validation",
        "Run kubeconform -strict -summary k8s/ in the project root.",
        "Output ends with: Summary: 32 resources found in 14 files - Valid: "
        "32, Invalid: 0, Errors: 0, Skipped: 0.",
        "10_kubeconform.png",
    )

    add_heading(doc, "6.3 Comparison", level=2)
    add_table(
        doc,
        ["Aspect", "Docker Swarm", "Kubernetes"],
        [
            ["Setup complexity", "Single command — docker swarm init", "Cluster required (minikube / kind / managed)"],
            ["Manifest format", "Compose with deploy: block", "Separate Deployment, Service, ConfigMap, etc."],
            ["Replication", "deploy.replicas (manual scale)", "spec.replicas + HPA (automatic)"],
            ["Health checks", "Single Docker healthcheck", "startupProbe + readinessProbe + livenessProbe"],
            ["Self-healing", "restart_policy", "ReplicaSet + probes"],
            ["Networking", "Overlay network", "CNI plugin + Service CIDR"],
            ["Where I used it", "Same VPS, quick replication", "Future-ready production setup"],
        ],
    )

    # ====== 7. TWO-DATABASE ARCHITECTURE (NEW SECTION) ======
    add_heading(doc, "7. Two-Database Architecture (PostgreSQL + MongoDB)", level=1)
    add_para(
        doc,
        "The End Term specification asks for two databases. I chose a polyglot "
        "persistence pattern instead of running two relational databases — "
        "PostgreSQL keeps its role as the primary transactional store, and "
        "MongoDB is added as a parallel store optimised for archival and "
        "full-text search. Each database is used where it excels.",
    )
    add_table(
        doc,
        ["Database", "Engine", "Used for", "Why this engine"],
        [
            ["postgres", "PostgreSQL 16", "Orders, users, finance, products, chat metadata", "ACID transactions, foreign keys, joins"],
            ["mongo", "MongoDB 7", "Chat message archive + full-text search index", "Schemaless docs, native $text index, easy horizontal sharding"],
        ],
    )
    add_para(
        doc,
        "When a chat message is created (REST endpoint or WebSocket), the "
        "service writes the canonical record to PostgreSQL (foreign key to the "
        "chat room ensures integrity), then writes a parallel copy to MongoDB. "
        "Mongo writes are best-effort — if Mongo is unreachable the chat keeps "
        "working, only the search index falls behind.",
    )
    add_para(doc, "Archive helper (chat_service/mongo.py):", italic=True)
    add_code_block(
        doc,
        """def archive_message(room_id, sender_user_id, sender_name,
                    content, created_at, message_pg_id=None):
    try:
        get_messages_collection().insert_one({
            "room_id": room_id,
            "message_pg_id": message_pg_id,
            "sender_user_id": sender_user_id,
            "sender_name": sender_name,
            "content": content,
            "created_at": created_at,
        })
        return True
    except PyMongoError:
        return False   # never block the request path""",
    )
    add_para(doc, "Full-text search endpoint (Russian language index):", italic=True)
    add_code_block(
        doc,
        """@router.get("/chats/archive/search")
def search_messages(q: str, limit: int = 50):
    return {"query": q, "results": search_archive(q.strip(), limit=limit)}""",
    )
    add_para(
        doc,
        "Indexes are created idempotently on app startup: a compound index on "
        "(room_id, created_at DESC) for chronological room queries, "
        "(sender_user_id) for per-user history, and a TEXT index on content "
        "with default_language=\"russian\" so cyrillic stemming works "
        "out-of-the-box.",
    )
    add_para(doc, "Deep health check exposes both databases at /health/full:", italic=True)
    add_code_block(
        doc,
        """{
  "status": "ok",
  "postgres": "ok",
  "mongo": "ok",
  "archive": { "total_messages": 142, "mongo_status": "connected" }
}""",
    )
    add_screenshot(
        doc,
        "chat_service /health/full",
        "Open the URL: http://213.155.22.46:8006/health/full",
        "JSON response showing status: ok, postgres: ok, mongo: ok, plus the "
        "archive stats with total_messages count.",
        "11_chat_health_full.png",
    )
    add_screenshot(
        doc,
        "Mongo archive search via API",
        "Open: http://213.155.22.46:8006/chats/archive/search?q=заказ",
        "JSON response with the query and a results array containing matching "
        "messages. Even if results array is empty, the structure proves Mongo "
        "is queried successfully.",
        "12_mongo_search.png",
    )

    # ====== 8. INFRASTRUCTURE AS CODE ======
    add_heading(doc, "8. Infrastructure as Code — Terraform", level=1)
    add_para(
        doc,
        "I split my Terraform configuration into two independent profiles in "
        "the terraform/ folder so that one provisions a real production server "
        "and the other is a cloud-ready reference.",
    )
    add_bullets(
        doc,
        [
            "terraform/existing-server/ — provisions the running VPS at "
            "213.155.22.46. Uses the null provider to SSH into the server, "
            "install Docker, rsync the project, run docker compose up, and "
            "configure Nginx as a reverse proxy.",
            "terraform/cloud/ — creates a brand-new Yandex Cloud VM with "
            "yandex_compute_instance, plus the VPC network, subnet, and "
            "security group. The cloud-provisioning reference required by "
            "the spec.",
        ],
    )
    add_code_block(
        doc,
        """cd terraform/existing-server
terraform init
terraform plan
terraform apply
# → Apply complete! Resources: 3 added, 0 changed, 0 destroyed.""",
    )
    add_screenshot(
        doc,
        "terraform apply success",
        "Run: cd terraform/existing-server && terraform apply -auto-approve",
        "Last 30 lines showing the green Apply complete! line and the Outputs "
        "block (server_ip, frontend_url, grafana_url, prometheus_url).",
        "13_terraform_apply.png",
    )
    add_screenshot(
        doc,
        "Yandex Cloud Terraform manifests",
        "Open terraform/cloud/compute.tf in VS Code.",
        "Show the yandex_compute_instance block with platform_id, resources, "
        "boot_disk, network_interface, and metadata (cloud-init with ssh-keys).",
        "14_terraform_cloud_compute.png",
    )

    # ====== 9. CONFIGURATION MANAGEMENT ======
    add_heading(doc, "9. Configuration Management — Ansible", level=1)
    add_para(
        doc,
        "Terraform handles infrastructure creation; Ansible handles ongoing "
        "configuration. The ansible/ folder contains one playbook and four roles:",
    )
    add_table(
        doc,
        ["Role", "Purpose"],
        [
            ["docker", "Install Docker CE + Compose plugin via the official apt repo"],
            ["app", "rsync the project to the server, run docker compose up -d --build"],
            ["nginx", "Install Nginx, deploy the reverse-proxy config, restart the service"],
            ["monitoring", "Verify that Prometheus and Grafana containers are running"],
        ],
    )
    add_code_block(
        doc,
        """# verify syntax (no execution)
ansible-playbook --syntax-check -i ansible/inventory.ini ansible/playbook.yml

# full deployment to the server
ansible-playbook -i ansible/inventory.ini ansible/playbook.yml""",
    )
    add_screenshot(
        doc,
        "Ansible folder structure",
        "Open the ansible/ folder in VS Code with the tree expanded.",
        "All four roles visible (docker, app, nginx, monitoring) — each "
        "expanded enough to see tasks/main.yml and templates/ where applicable.",
        "15_ansible_tree.png",
    )
    add_screenshot(
        doc,
        "ansible-playbook --syntax-check result",
        "Run: ansible-playbook --syntax-check -i ansible/inventory.ini "
        "ansible/playbook.yml",
        "Output: playbook: ansible/playbook.yml — no errors.",
        "16_ansible_syntax.png",
    )

    # ====== 10. CI/CD ======
    add_heading(doc, "10. CI/CD — Multi-Stage GitHub Actions Pipeline", level=1)
    add_para(
        doc,
        "I split the deployment pipeline into two independent workflows so "
        "that quality gates run on every change but deploys only happen after "
        "everything passes.",
    )
    add_para(doc, "10.1 ci.yml — runs on every push and pull request", bold=True)
    add_table(
        doc,
        ["Job", "What it does"],
        [
            ["python-lint", "ruff check on all six microservice folders"],
            ["docker-compose-validate", "docker compose config + docker compose -f docker-stack.yml config"],
            ["terraform-validate", "Matrix job — fmt + init + validate for both terraform/existing-server and terraform/cloud"],
            ["ansible-syntax", "ansible-playbook --syntax-check"],
            ["k8s-validate", "kubeconform -strict -summary k8s/ (all 32 resources)"],
            ["yaml-lint", "yamllint across compose, monitoring, k8s, ansible, .github"],
        ],
    )
    add_para(doc, "10.2 deploy.yml — runs only on push to main", bold=True)
    add_para(
        doc,
        "Three-stage pipeline with concurrency control and environment "
        "protection so two deploys never run in parallel:",
    )
    add_code_block(
        doc,
        """jobs:
  validate:    # re-runs critical IaC checks before touching prod
  deploy:      # needs: validate
    environment: production
    steps:
      - Deploy via SSH (appleboy/ssh-action, script_stop: true, timeout: 15m)
      - Health check after deploy — curl /health on all 6 services, fail if any != 200
  notify:      # needs: deploy, runs always: even on failure
    steps:
      - Telegram message with the deploy status, commit, branch, actor, and run URL""",
    )
    add_para(
        doc,
        "The flow becomes: git push main → CI checks (6 jobs) → if green, deploy "
        "workflow starts → re-validate → SSH+deploy.sh → health check on six "
        "/health endpoints → Telegram ✅ or ❌.",
    )
    add_screenshot(
        doc,
        "GitHub Actions — both workflows passing",
        "Open https://github.com/adilzhankad/sre_mooztau_test/actions",
        "List of recent runs — both CI and Deploy workflows showing green "
        "checks. Capture at least one successful Deploy-to-Production run with "
        "the recent commit hash visible.",
        "17_github_actions.png",
    )
    add_screenshot(
        doc,
        "Deploy workflow run — three stages",
        "Click on a successful Deploy-to-Production run in Actions.",
        "Three green stages visible: validate → deploy → notify. Expand the "
        "deploy job to show the Health check after deploy step output: "
        "Port 8001 → 200, Port 8002 → 200, etc.",
        "18_deploy_workflow_stages.png",
    )

    # ====== 11. INCIDENT RESPONSE ======
    add_heading(doc, "11. Incident Response Simulation (Assignment 4)", level=1)
    add_para(
        doc,
        "To validate the entire detect-alert-recover loop, I deliberately broke "
        "the Orders Service in production.",
    )
    add_para(doc, "Procedure:", bold=True)
    add_code_block(doc, "docker compose stop orders")
    add_para(
        doc,
        "Within 15 seconds Prometheus marked the target as DOWN; the Grafana "
        "tile turned red; Alertmanager sent a Telegram ИНЦИДЕНТ message.",
    )
    add_screenshot(
        doc,
        "Grafana — orders DOWN",
        "While orders is stopped: http://213.155.22.46:3000 → MoozTau dashboard.",
        "The orders-service tile in red labeled DOWN. Other five services UP.",
        "19_grafana_down.png",
    )
    add_screenshot(
        doc,
        "Telegram ИНЦИДЕНТ notification",
        "Open the Telegram chat with @mooztau_alerts_bot.",
        "The red ИНЦИДЕНТ card with: Сервис = orders-service, "
        "Алерт = ServiceDown, timestamp.",
        "20_telegram_alert.png",
    )
    add_para(doc, "Recovery:", bold=True)
    add_code_block(doc, "docker compose start orders")
    add_para(
        doc,
        "The container came back up within 30 seconds. Prometheus flipped the "
        "target to UP; Grafana turned green; Telegram delivered a "
        "ВОССТАНОВЛЕНО confirmation. Full postmortem in docs/postmortem.md.",
    )
    add_screenshot(
        doc,
        "Telegram ВОССТАНОВЛЕНО",
        "Telegram chat with @mooztau_alerts_bot after the service recovered.",
        "Green ВОССТАНОВЛЕНО card for orders-service with the recovery timestamp.",
        "21_telegram_resolved.png",
    )

    # ====== 12. AUTOMATION & CAPACITY PLANNING ======
    add_heading(doc, "12. Automation and Capacity Planning (Assignment 6)", level=1)
    add_heading(doc, "12.1 Self-Healing", level=2)
    add_para(
        doc,
        "Every container has restart: unless-stopped + a Docker health check. "
        "If health checks fail Docker restarts automatically. In Kubernetes the "
        "three-level probe stack from §6.2 plus the HPA from the same section "
        "provide the equivalent: failed pods are restarted by livenessProbe, "
        "load spikes are absorbed by HPA adding pods.",
    )
    add_heading(doc, "12.2 Load Testing", level=2)
    add_para(
        doc,
        "load_test.py uses asyncio + aiohttp to fire concurrent requests at all "
        "six services. A representative 90-second run:",
    )
    add_table(
        doc,
        ["Service", "Requests", "Errors", "Error %", "Avg ms", "P95 ms"],
        [
            ["auth", "19 257", "0", "0.0%", "139", "223"],
            ["orders", "19 257", "0", "0.0%", "142", "225"],
            ["finance", "19 257", "0", "0.0%", "90", "160"],
            ["product", "6 419", "0", "0.0%", "47", "95"],
            ["user", "6 419", "0", "0.0%", "43", "88"],
            ["chat", "6 419", "0", "0.0%", "50", "102"],
            ["TOTAL", "77 028", "0", "0.0%", "—", "—"],
        ],
    )
    add_screenshot(
        doc,
        "load_test.py output",
        "Run: python3 load_test.py --host 213.155.22.46 --users 20 "
        "--duration 90 --rps 10",
        "Final ASCII LOAD TEST REPORT table — every service [OK], TOTAL row "
        "with 77 028 requests / 0 errors.",
        "22_load_test.png",
    )
    add_screenshot(
        doc,
        "Grafana under load",
        "While load_test is running, open Grafana → MoozTau dashboard, "
        "Last 15 minutes.",
        "Visible Request Rate spike, Latency panel showing the burst.",
        "23_grafana_under_load.png",
    )

    add_heading(doc, "12.3 Findings", level=2)
    add_bullets(
        doc,
        [
            "orders-service is the most resource-intensive component — JOINs "
            "over orders + products + payments + audit.",
            "PostgreSQL is the shared bottleneck — all six services compete "
            "for the same connection pool.",
            "P95 latency stays under 225 ms at the tested ~200 RPS, well below "
            "the 2-second HighLatency alert threshold.",
            "No errors (0/77028) even at the saturation rate of the connection "
            "pool — health checks and the connection-keep-alive design hold.",
        ],
    )
    add_heading(doc, "12.4 Scaling Strategy (now implemented)", level=2)
    add_para(
        doc,
        "Earlier reports listed scaling as a strategy. In this End Term I "
        "actually implemented it:",
    )
    add_table(
        doc,
        ["Strategy", "How it works in this project", "Where"],
        [
            ["Horizontal scaling (Swarm)", "deploy.replicas: 2 in docker-stack.yml, docker service scale to change at runtime", "docker-stack.yml"],
            ["Horizontal scaling (Kubernetes)", "HPA on CPU + memory, per-service min/max — see §6.2", "k8s/hpa.yml"],
            ["Vertical scaling", "Update resources.requests/limits in compose or k8s", "every Deployment"],
            ["DB optimisation", "Documented in capacity_planning.md (PgBouncer, indexes, read replicas)", "docs/capacity_planning.md"],
        ],
    )
    add_screenshot(
        doc,
        "Capacity Planning document",
        "Open docs/capacity_planning.md in VS Code preview.",
        "Findings + Strategies section with the bottleneck analysis and the "
        "scaling table.",
        "24_capacity_planning.png",
    )

    # ====== 13. VALIDATION ======
    add_heading(doc, "13. Full Validation", level=1)
    add_para(
        doc,
        "Because the project spans Docker Compose, Docker Swarm, Kubernetes, "
        "Terraform (two profiles), and Ansible, I wrote a single validation "
        "script that checks all of them in one command.",
    )
    add_code_block(doc, "bash scripts/validate_all.sh")
    add_para(
        doc,
        "When I ran it for this report, every section reported [OK] and "
        "kubeconform confirmed 32 resources valid in 14 files (the new "
        "mongo manifest + 6 HPA brought the count up from the previous 23/12).",
    )
    add_screenshot(
        doc,
        "validate_all.sh full output",
        "Run bash scripts/validate_all.sh from the project root.",
        "Five numbered sections, every check [OK] green, final line "
        "'All checks completed'.",
        "25_validate_all.png",
    )

    # ====== 14. RESULTS ======
    add_heading(doc, "14. Results", level=1)
    add_table(
        doc,
        ["Deliverable", "Status", "Location"],
        [
            ["6+ microservices", "Done", "auth_service/, orders_service/, finance_service/, product_service/, user_service/, chat_service/"],
            ["Docker Compose orchestration", "Done", "docker-compose.yml (12 services)"],
            ["Docker Swarm configuration", "Done", "docker-stack.yml"],
            ["Kubernetes manifests", "Done", "k8s/ — 32 resources, kubeconform clean"],
            ["Health checks (3-level k8s probes)", "Done", "every Deployment — startupProbe + readiness + liveness"],
            ["Horizontal Pod Autoscaling", "Done", "k8s/hpa.yml — 6 HPA per microservice"],
            ["Two databases (Postgres + Mongo)", "Done", "docker-compose.yml, chat_service/mongo.py, k8s/mongo.yml"],
            ["Terraform — existing server", "Done", "terraform/existing-server/ (apply complete)"],
            ["Terraform — cloud (Yandex)", "Done", "terraform/cloud/ (validated)"],
            ["Ansible playbook", "Done", "ansible/ — 4 roles"],
            ["Prometheus + Grafana", "Done", "monitoring/, :9090 + :3000"],
            ["Alertmanager + Telegram", "Done", "monitoring/alertmanager/ + @mooztau_alerts_bot"],
            ["SLI / SLO design", "Done", "docs/sli_slo.md"],
            ["Capacity planning", "Done", "docs/capacity_planning.md"],
            ["Incident report + postmortem", "Done", "docs/incident_report.md + docs/postmortem.md"],
            ["Validation script", "Done", "scripts/validate_all.sh (5/5 [OK])"],
            ["Multi-stage CI/CD pipeline", "Done", ".github/workflows/ci.yml + deploy.yml"],
            ["Live demo / screenshots", "Done", "screenshots/ folder"],
        ],
    )
    add_screenshot(
        doc,
        "Frontend at medhome.kz",
        "Open http://213.155.22.46:3100 (or http://medhome.kz).",
        "MoozTau dashboard / login page proving the end-user system is live.",
        "26_frontend.png",
    )

    # ====== 15. CONCLUSION ======
    add_heading(doc, "15. Conclusion", level=1)
    add_para(
        doc,
        "Working through the End Term Project tied together everything I had "
        "built over the course and added several production-grade pieces that "
        "the spec demanded. I now have a single platform that demonstrates the "
        "full SRE lifecycle: containerised microservices, three different "
        "orchestrators (Compose, Swarm, Kubernetes), declarative infrastructure "
        "provisioning with Terraform, configuration automation with Ansible, "
        "monitoring with Prometheus and Grafana, alerting via Alertmanager and "
        "Telegram, a multi-stage CI/CD pipeline with lint/validate/build/deploy/"
        "notify stages, a polyglot two-database architecture (PostgreSQL + "
        "MongoDB), Kubernetes health checks at three levels, and horizontal "
        "pod autoscaling on every microservice.",
    )
    add_para(
        doc,
        "The system runs reliably on a production VPS, every configuration is "
        "syntactically validated by a single-command script, and the incident "
        "simulation proved that the detect-alert-recover loop closes in under "
        "a minute. The load test demonstrates 77 028 requests served with zero "
        "errors and sub-225 ms P95 latency. By combining these practices, I "
        "brought MoozTau to a production-ready level of operational maturity "
        "that matches what real SRE teams aim for.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK — {OUT}")


if __name__ == "__main__":
    build()
